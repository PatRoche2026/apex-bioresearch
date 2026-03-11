"""Generate HW4 writeup as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# Heading 1 style
h1_style = doc.styles["Heading 1"]
h1_style.font.name = "Calibri"
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
h1_style.paragraph_format.space_before = Pt(0)
h1_style.paragraph_format.space_after = Pt(6)

# Heading 2 style
h2_style = doc.styles["Heading 2"]
h2_style.font.name = "Calibri"
h2_style.font.size = Pt(11)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x4A)
h2_style.paragraph_format.space_before = Pt(10)
h2_style.paragraph_format.space_after = Pt(4)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Multi-Agent Orchestration Infrastructure\nfor Biotech R&D Decision-Making")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
title.paragraph_format.space_after = Pt(4)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Patrick (Pat) Ovando Roche, PhD\nMAS.664 — AI Studio: AI Agents and Agentic Web | HW4 | March 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
author.paragraph_format.space_after = Pt(2)

# Thin line
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run("─" * 80)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
line.paragraph_format.space_after = Pt(8)

# =====================================================================
# SECTION 1
# =====================================================================
doc.add_heading("1. Topic & Why It Matters", level=2)

doc.add_paragraph(
    "Over twelve years working as a scientist across big pharma and biotech — from early "
    "discovery through manufacturing transfer — I have watched promising therapeutics fail "
    "for reasons that had nothing to do with the underlying biology. Targets selected from "
    "a single academic publication rather than systematic, data-driven evaluation. Promising "
    "programs killed by bureaucratic friction between departments: discovery teams, safety "
    "groups, and clinical operations each applying different mental models, unable to "
    "reconcile conflicting evidence, stalled by ego and institutional inertia. These are "
    "human coordination failures, not scientific ones, and they are responsible for more "
    "drug attrition than any single technical challenge."
)

doc.add_paragraph(
    "Multi-agent orchestration infrastructure offers a path forward. Instead of one "
    "scientist trying to hold genetics, clinical trials, single-cell biology, safety "
    "pharmacology, and competitive intelligence in their head simultaneously, specialized "
    "AI agents — each with deep access to domain-specific databases via MCP and autonomous "
    "tool use — can evaluate targets from multiple lenses in parallel, debate each other's "
    "conclusions, and surface disagreements explicitly. A recent preprint from Stanford "
    "(Swanson et al., bioRxiv, Feb 2026) demonstrated this with the \"Virtual Biotech\": "
    "11 agents, 10 MCP servers, 100+ tools that processed 55,984 clinical trials. I set "
    "out to build my own version — APEX — and go further by adding agent debate, "
    "RAG-powered institutional knowledge, and human-in-the-loop CEO feedback."
)

# =====================================================================
# SECTION 2
# =====================================================================
doc.add_heading("2. Existing Landscape", level=2)

p = doc.add_paragraph()
run = p.add_run("Orchestration frameworks. ")
run.bold = True
p.add_run(
    "LangGraph (explicit stateful graphs with parallel fan-out/fan-in — what I used), "
    "CrewAI (role-based, simpler but less control over state), AutoGen (Microsoft, "
    "multi-agent conversations), Claude Agent SDK (Anthropic, used by Stanford Virtual Biotech)."
)

p = doc.add_paragraph()
run = p.add_run("Protocols. ")
run.bold = True
p.add_run(
    "MCP — Model Context Protocol (Anthropic) standardizes how agents access tools and "
    "data sources. A2A — Agent-to-Agent (Google) enables inter-agent communication. "
    "Both are emerging standards shaping the agentic infrastructure layer."
)

p = doc.add_paragraph()
run = p.add_run("Notable biotech agent projects. ")
run.bold = True
p.add_run(
    "Stanford Virtual Biotech (Swanson et al., Feb 2026): 11 agents across 4 scientific "
    "divisions, 10 MCP servers, processed 55,984 clinical trials — the closest comparable "
    "to APEX. Prompt-to-Pill (Oxford, Dec 2025): modular drug discovery pipeline. "
    "PharmaSwarm: hypothesis-driven agent swarm. BioMNI: general biomedical AI agent. "
    "CellVoyager: autonomous computational biology."
)

p = doc.add_paragraph()
run = p.add_run("Course concepts applied. ")
run.bold = True
p.add_run(
    "Join39 (course agent registry/platform), NANDA (decentralized agent architecture "
    "from course lectures), and the infrastructure layers from L5-1: AutoML, PrivateML, "
    "VerifiableAI, KnowledgePricing, Orchestration, Population Models, CrowdUX."
)

# =====================================================================
# SECTION 3
# =====================================================================
doc.add_heading("3. What I Built & Tested (APEX)", level=2)

doc.add_paragraph(
    "I built APEX (Agentic Pipeline for Executive Decisions) — a multi-agent biotech "
    "executive debate system deployed end-to-end on Railway with a React frontend. "
    "APEX goes beyond the Stanford Virtual Biotech by adding structured inter-agent "
    "debate, RAG-powered institutional knowledge, and human-in-the-loop CEO feedback."
)

p = doc.add_paragraph()
run = p.add_run("Architecture. ")
run.bold = True
p.add_run(
    "A 10-node LangGraph StateGraph with parallel fan-out/fan-in and conditional looping. "
    "Five AI agents with named personas and devil's advocate triggers:"
)

agents = [
    ("Dr. Elena Vasquez (CSO)", "evaluates target biology, genetic evidence, mechanism of action"),
    ("Dr. Marcus Wei (CTO)", "evaluates druggability, modality selection, manufacturing feasibility"),
    ("Dr. Sarah Okonkwo (CMO)", "evaluates clinical landscape, regulatory path, trial design"),
    ("Dr. James Harrington (CBO)", "evaluates market size, competitive dynamics, IP position"),
    ("Dr. Amara Chen (Portfolio Director)", "synthesizes all assessments into weighted GO/CONDITIONAL GO/NO-GO verdict"),
]
for name, desc in agents:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name} — ")
    run.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
run = p.add_run("Pipeline flow. ")
run.bold = True
p.add_run(
    "Research Scout (PubMed multi-query search with LLM relevance filtering) → 4 parallel "
    "executive assessments → debate round (each agent reads all 4 assessments and challenges "
    "the weakest argument) → Portfolio Director synthesis with weighted scoring (Scientific "
    "0.30, Technical 0.25, Clinical 0.25, Commercial 0.20). Conditional loop: if confidence "
    "< 60%, agents debate again (max 2 rounds). Rate-limited with asyncio.Semaphore(3) and "
    "60-second per-node timeouts."
)

p = doc.add_paragraph()
run = p.add_run("Domain-specific autonomous tools (9 total). ")
run.bold = True
p.add_run(
    "Each agent has different tools matching their expertise, enforced server-side. "
    "CSO: PubMed + UniProt + STRING-DB + PMC Full Text. CTO: PubMed + ChEMBL + Open Targets "
    "Tractability + PMC Full Text. CMO: ClinicalTrials.gov + PubMed + Open Targets Safety + "
    "PMC Full Text. CBO: PubMed + Open Targets + ClinicalTrials.gov + PMC Full Text. The "
    "search_and_read_papers tool retrieves full-text articles from PubMed Central (PMC) via "
    "Entrez.elink ID mapping and XML body parsing, truncated to 3,000 characters. Agents "
    "autonomously request tools during assessment via a two-pass approach: first LLM call "
    "produces assessment with optional TOOL_REQUEST lines → tools are executed → second LLM "
    "call revises assessment with tool results."
)

p = doc.add_paragraph()
run = p.add_run("RAG knowledge bases. ")
run.bold = True
p.add_run(
    "ChromaDB vector store with Voyage AI embeddings (voyage-3-lite). 382 document chunks "
    "from 53 real PMC papers across 5 persistent collections — CSO: 130 chunks from 15 papers "
    "(GWAS, Mendelian randomization, CRISPR functional genomics); CTO: 65 chunks from 12 "
    "papers (druggability, ADC design, BBB delivery); CMO: 109 chunks from 15 papers "
    "(adaptive trials, FDA breakthrough therapy); CBO: 29 chunks from 5 papers (orphan drug "
    "pricing, market access); Shared: 49 chunks from 6 papers (drug discovery, OSMR/IBD). "
    "Papers were downloaded via a curated PubMed search script with PMC full-text enrichment "
    "(up to 8,000 characters per paper). Each agent retrieves role-specific + shared "
    "institutional knowledge before reasoning."
)

p = doc.add_paragraph()
run = p.add_run("Real results — OSMR/ulcerative colitis. ")
run.bold = True
p.add_run(
    "The pipeline found 8 PubMed papers including the seminal West et al. 2017 Nature "
    "Medicine study (PMID: 28368383). CSO rated scientific validity 6/10 citing lack of "
    "independent GWAS replication. CTO rated technical feasibility 7/10 noting antibody "
    "modality is viable but gp130 family selectivity is a risk. CMO flagged the terminated "
    "MOONGLOW Phase II trial (NCT03145246) and proposed biomarker-guided enrollment. CBO "
    "scored commercial potential 4/10 due to competitive anti-TNF/anti-IL-23 alternatives. "
    "Portfolio Director verdict: CONDITIONAL GO at 75% confidence with 5.6/10 weighted score."
)

p = doc.add_paragraph()
run = p.add_run("Three end-to-end test queries. ")
run.bold = True

# Table for results
table = doc.add_table(rows=4, cols=5)
table.style = "Light Shading Accent 1"
headers = ["Query", "Verdict", "Confidence", "Weighted Score", "Rounds"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

data = [
    ["OSMR / Ulcerative Colitis", "CONDITIONAL GO", "75%", "5.6/10", "1"],
    ["GLP-1R / Alzheimer's Disease", "NO-GO", "75%", "4.8/10", "2*"],
    ["CD47 / Cancer Immunotherapy", "NO-GO", "85%", "4.2/10", "1"],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("*Triggered conditional loop — confidence was < 60% after round 1.").font.size = Pt(8)
p.runs[0].font.italic = True

p = doc.add_paragraph()
run = p.add_run("Human-in-the-loop CEO feedback. ")
run.bold = True
p.add_run(
    "After any evaluation, a CEO can submit feedback via POST /feedback/{session_id} or "
    "WebSocket /ws/feedback/{session_id}. The feedback is injected as a \"CEO DIRECTIVE\" "
    "into all agent prompts, and the full pipeline re-runs with scout data carried forward "
    "(avoiding redundant PubMed calls). Feedback history is tracked across rounds."
)

p = doc.add_paragraph()
run = p.add_run("Deployment. ")
run.bold = True
p.add_run(
    "FastAPI backend with WebSocket streaming on Railway (Docker, Dockerfile with "
    "python:3.11-slim). React frontend built in Lovable showing real-time agent debate "
    "with persona cards, score visualizations, and markdown rendering. "
    "Join39 app registered as \"APEX BioResearch Agent.\""
)

# =====================================================================
# SECTION 4
# =====================================================================
doc.add_heading("4. What I Learned", level=2)

p = doc.add_paragraph()
run = p.add_run("What worked. ")
run.bold = True
p.add_run(
    "LangGraph's explicit graph definition with fan-out/fan-in makes parallel agent "
    "execution transparent and debuggable — you can see exactly which node ran, what "
    "state was passed, and where failures occur. The Annotated[list, operator.add] "
    "reducer pattern elegantly handles parallel state merges without data loss. Voyage AI "
    "embeddings are a lightweight alternative to local sentence-transformers (~50KB vs "
    "~2GB PyTorch), enabling RAG on Railway without blowing up container size. The "
    "two-pass tool use approach (agent requests tools → execute → agent revises with "
    "results) is more reliable than direct LangChain tool binding because the agent "
    "controls when and why it needs external data."
)

p = doc.add_paragraph()
run = p.add_run("What was confusing. ")
run.bold = True
p.add_run(
    "LangGraph's state management requires careful upfront schema design — changing the "
    "TypedDict schema after building nodes requires refactoring all of them. The "
    "executive_scores field initially crashed parallel execution because four agents "
    "writing to the same dict key overwrote each other; this required a custom "
    "_merge_dicts reducer. MCP documentation assumes familiarity with SSE transport "
    "that isn't intuitive. Windows encoding issues (cp1252 vs UTF-8) caused crashes "
    "when agents output Unicode characters in their assessments."
)

p = doc.add_paragraph()
run = p.add_run("Limitations and tradeoffs. ")
run.bold = True
p.add_run(
    "ClinicalTrials.gov API returns 403 Forbidden from non-US IPs locally (works on "
    "Railway's US-based servers). Each full evaluation costs ~$0.25 in Anthropic API credits "
    "(Claude Sonnet for 10 LLM calls) — tracked per-node via usage_metadata token counting "
    "and aggregated in all API/WebSocket responses as estimated_cost_usd. No built-in "
    "agent trust or verification mechanism — when agents disagree, the Portfolio "
    "Director's resolution is only as good as its prompt engineering. Session storage is "
    "in-memory (resets on Railway redeploy); production would need a database."
)

# =====================================================================
# SECTION 5
# =====================================================================
doc.add_heading("5. Why This Interests Me", level=2)

doc.add_paragraph(
    "I am building APEX as my Demo Day project for MAS.664 — an Agentic Pipeline for "
    "Executive Decisions that compresses months of biotech investment committee "
    "deliberation into minutes. The system addresses a real problem I've seen repeatedly "
    "in drug discovery: promising targets killed not by bad science but by departmental "
    "silos, ego, and information loss between functional experts. APEX replaces that with "
    "AI agents that debate transparently, cite evidence, and produce auditable reasoning "
    "chains."
)

doc.add_paragraph(
    "My next steps: voice interaction via ElevenLabs TTS and Whisper STT so users can "
    "verbally challenge the verdict, video avatars via D-ID for each executive persona, "
    "and integration with additional biomedical databases — CELLxGENE Census for "
    "single-cell expression, GEO for transcriptomics, and GWAS Catalog for genetic "
    "associations — to create the most comprehensive AI-driven drug target evaluation "
    "platform in the class."
)

# =====================================================================
# FOOTER
# =====================================================================
line2 = doc.add_paragraph()
line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line2.add_run("─" * 80)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
line2.paragraph_format.space_before = Pt(10)
line2.paragraph_format.space_after = Pt(6)

footer_items = [
    ("Join39 App: ", "APEX BioResearch Agent"),
    ("Backend: ", "https://independent-benevolence-production-46d3.up.railway.app"),
    ("API Docs: ", "https://independent-benevolence-production-46d3.up.railway.app/docs"),
    ("GitHub: ", "https://github.com/PatRoche2026/apex-bioresearch"),
    ("Framework: ", "LangGraph + FastAPI + ChromaDB + Voyage AI + Claude Sonnet + 53 PMC Papers"),
]

for label, value in footer_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run = p.add_run(value)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

# Save
output_path = "HW4_Patrick_Roche_FINAL.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
